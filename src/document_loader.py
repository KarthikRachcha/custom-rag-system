import os
from typing import List, Dict, Any
from pathlib import Path
from pypdf import PdfReader

class DocumentLoader:
    def __init__(self, data_dir: str):
        self.data_dir = Path(data_dir)
        self.supported_extensions ={'.txt', '.pdf','.docx'}

    def load_documents(self) -> List[Dict[str,Any]]:
        """Load all supported documents from data directory"""
        documents=[]

        for file_path in self.data_dir.rglob('*'):
            if file_path.suffix.lower() in self.supported_extensions:
                doc = self._load_single_document(file_path)
                if doc:
                    documents.append(doc)
        
        return documents
    
    def _load_single_document(self, file_path: Path) -> Dict[str, Any]:
        """Load a single document based on its extension"""
        try:
            if file_path.suffix.lower() == '.txt':
                return self._load_text_file(file_path)
            elif file_path.suffix.lower() == '.pdf':
                return self._load_pdf_file(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._load_docx_file(file_path)
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
            return None
    
    def _load_text_file(self, file_path: Path) -> Dict[str, Any]:
        """Load a text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            'content': content,
            'source': str(file_path),
            'type': 'text',
            'metadata': {
                'filename': file_path.name,
                'size': len(content)
            }
        }
    
    def _load_pdf_file(self, file_path: Path) -> Dict[str, Any]:
        """Load a PDF file using pypdf"""
        try:
            reader = PdfReader(file_path)
            content = ""
            
            # Extract text from all pages
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            if not content.strip():
                print(f"Warning: No text extracted from PDF {file_path}")
                return None
            
            return {
                'content': content.strip(),
                'source': str(file_path),
                'type': 'pdf',
                'metadata': {
                    'filename': file_path.name,
                    'pages': len(reader.pages),
                    'size': len(content)
                }
            }
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")
            return None
    
    def _load_docx_file(self, file_path: Path) -> Dict[str, Any]:
        """Load a Word document using python-docx"""
        try:
            from docx import Document
            doc = Document(file_path)
            content = ""
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            if not content.strip():
                print(f"Warning: No text extracted from DOCX {file_path}")
                return None
            
            return {
                'content': content.strip(),
                'source': str(file_path),
                'type': 'docx',
                'metadata': {
                    'filename': file_path.name,
                    'paragraphs': len(doc.paragraphs),
                    'size': len(content)
                }
            }
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")
            return None